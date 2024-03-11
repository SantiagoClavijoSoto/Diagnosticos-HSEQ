from django.shortcuts import render
from django.http.response import Http404
from rest_framework.views import APIView
from rest_framework.response import Response
from django.http import HttpResponse
from django.http.response import JsonResponse
from rest_framework.exceptions import ValidationError
from .models import Fase1, Fase2, Fase3, Fase4
from .serializers import FasesSerializer
from docx.shared import Pt
from docx import Document


class FasesView(APIView):
    serializer_class = FasesSerializer

    def get_model(self, fase_number):
        model_mapping = {1: Fase1, 2: Fase2, 3: Fase3, 4: Fase4}
        selected_model = model_mapping.get(fase_number)
        return selected_model

    def get_object(self, model, pk):
        try:
            return model.objects.get(id=pk)
        except model.DoesNotExist:
            raise Http404

    def get(self, request, fase_number, pk=None):
        selected_model = self.get_model(fase_number)

        if selected_model is None:
            return JsonResponse("Número de fase no válido.", status=400, safe=False)

        if pk:
            data = self.get_object(selected_model, pk)
            serializer = FasesSerializer(data)
        else:
            data = selected_model.objects.all()
            serializer = FasesSerializer(data, many=True)

        return Response(serializer.data)

    def post(self, request, fase_number):
        selected_model = self.get_model(fase_number)

        if selected_model is None:
            return JsonResponse("Número de fase no válido.", status=400, safe=False)

        data = request.data
        serializer = self.serializer_class(data=data, context={"model": selected_model})

        if serializer.is_valid():
            serializer.save()
            return JsonResponse("Ítem agregado exitosamente.", safe=False)

        return JsonResponse("No se pudo agregar ítem.", status=400, safe=False)

    def put(self, request, fase_number, pk=None):
        selected_model = self.get_model(fase_number)

        if selected_model is None:
            return JsonResponse("Número de fase no válido.", status=400, safe=False)

        fase_to_update = self.get_object(selected_model, pk)
        serializer = FasesSerializer(
            instance=fase_to_update,
            data=request.data,
            partial=True,
            context={"model": selected_model},
        )

        try:
            serializer.is_valid(raise_exception=True)
            serializer.save()
            return JsonResponse("Ítem actualizado con éxito.", safe=False)
        except ValidationError as e:
            return JsonResponse({"error": str(e)}, status=400)

    def delete(self, request, fase_number, pk):
        selected_model = self.get_model(fase_number)

        if selected_model is None:
            return JsonResponse("Número de fase no válido.", status=400, safe=False)

        fase_to_delete = self.get_object(selected_model, pk)
        fase_to_delete.delete()

        return JsonResponse("Ítem eliminado exitosamente.", safe=False)


def generar_informe(request):
    # Cargar la plantilla
    plantilla_path = "fases/documents/pesv.docx"
    # Crear un nuevo documento Word
    doc = Document(plantilla_path)

    paragraph = doc.add_paragraph()
    paragraph.space_after = Pt(29)  # Puedes ajustar el valor según tus necesidades
    paragraph.add_run("\t" * 5 + "ASISTENCIA TÉCNICA PLAN ESTRATÉGICO DE SEGURIDAD VIAL\nCONSULTORÍA Y CAPACITACIÓN HSEQ SAS").bold = True

    # # Agregar contenido al documento
    # doc.add_heading("ASISTENCIA TÉCNICA PLAN ESTRATÉGICO DE SEGURIDAD VIAL", level=1)
    # doc.add_heading("CONSULTORÍA Y CAPACITACIÓN HSEQ SAS", level=1)
    # doc.add_heading("CRONOGRAMA XXXXXX SECUENCIA XX", level=1)
    # doc.add_heading("ESTACIONAMIENTO LUGANO", level=1)
    # doc.add_heading("NIT 800529157-1", level=1)
    # doc.add_heading("BOGOTÁ", level=1)
    # doc.add_heading("AGOSTO DE 2023", level=1)

    paragraph = doc.add_paragraph("Realizado por:")
    doc.add_paragraph("ANGELA RUIZ ORTIZ")
    doc.add_paragraph("Consultor en Seguridad Vial")
    doc.add_paragraph("Licencia SST 5996 de 2015")

    # Guardar el documento en la respuesta HTTP
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = "attachment; filename=informe.docx"
    doc.save(response)

    return response
